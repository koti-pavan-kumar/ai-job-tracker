from fastapi import FastAPI, Depends, HTTPException, status, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from sqlalchemy.orm import Session
from typing import List
import jwt
from jwt.exceptions import PyJWTError as JWTError
import io
import os
from dotenv import load_dotenv
import bcrypt  # Using native python-bcrypt directly for Python 3.13 compatibility

import models
import schemas
import database
import scraper
import ai_service

# Document generation packages
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document

load_dotenv()
# Security Settings Configuration
SECRET_KEY = os.getenv("SECRET_KEY")
ALGORITHM = os.getenv("ALGORITHM", "HS256")

if not SECRET_KEY:
    raise RuntimeError("CRITICAL CRASH: SECRET_KEY configuration environment variable is missing!")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

def hash_password(password: str) -> str:
    # Encodes string to bytes, salts, and hashes natively
    salt = bcrypt.gensalt()
    return bcrypt.hashpw(password.encode('utf-8'), salt).decode('utf-8')

def verify_password(plain_password: str, hashed_password: str) -> bool:
    # Safely checks plane textual bytes against database strings
    return bcrypt.checkpw(plain_password.encode('utf-8'), hashed_password.encode('utf-8'))

# Initialize Database Schema Structures
models.Base.metadata.create_all(bind=database.engine)

app = FastAPI(title="AI Job Application Tracker API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173", 
        "https://ai-job-tracker-lljg.onrender.com"  # Paste your live frontend URL here
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/token")
def login(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(database.get_db)):
    # Look up the user dynamically inside your database instance
    user = db.query(models.User).filter(models.User.username == form_data.username).first()
    
    if not user or not verify_password(form_data.password, user.hashed_password):
        raise HTTPException(status_code=400, detail="Invalid workspace credential criteria values verified.")
        
    token = jwt.encode({"sub": form_data.username}, SECRET_KEY, algorithm=ALGORITHM)
    return {"access_token": token, "token_type": "bearer"}

def get_current_user(token: str = Depends(oauth2_scheme)):
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise HTTPException(status_code=401, detail="Token payload validation mapping collapsed.")
        return username
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid token parameters encountered.")

def get_user_job_or_raise(job_id: int, db: Session, username: str) -> models.JobApplication:
    """
    Secures the application against BOLA leaks by ensuring the resource 
    exists and belongs exclusively to the authenticated user.
    """
    # 1. Fetch the logged-in user's record
    user = db.query(models.User).filter(models.User.username == username).first()
    if not user:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid session context.")

    # 2. Query the job application matching both the job_id AND the user's id
    db_job = db.query(models.JobApplication).filter(
        models.JobApplication.id == job_id,
        models.JobApplication.user_id == user.id
    ).first()

    # 3. If it doesn't exist or doesn't belong to them, return a 404
    if not db_job:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, 
            detail="Requested job profile record not found in this workspace."
        )
        
    return db_job

@app.post("/jobs/analyze-url", response_model=schemas.JobCreate)
def analyze_url(request: schemas.URLIngestRequest):
    scraped_data = scraper.scrape_job_details(request.url)
    if not scraped_data or scraped_data.get("company_name") == "Error Parsing":
        raise HTTPException(status_code=400, detail="Unable to safely parse details from resource URL location.")
    
    skills = ai_service.analyze_and_parse_jd(scraped_data["raw_jd"])
    scraped_data["parsed_skills"] = skills
    return scraped_data

@app.post("/jobs", response_model=schemas.JobResponse)
def create_job(job: schemas.JobCreate, db: Session = Depends(database.get_db), current_username: str = Depends(get_current_user)):
    user = db.query(models.User).filter(models.User.username == current_username).first()
    
    db_job = models.JobApplication(**job.model_dump())
    if user:
        db_job.user_id = user.id # Link the job card to this user explicitly
        
    db.add(db_job)
    db.commit()
    db.refresh(db_job)
    return db_job

@app.get("/jobs", response_model=List[schemas.JobResponse])
def get_jobs(db: Session = Depends(database.get_db), current_username: str = Depends(get_current_user)):
    user = db.query(models.User).filter(models.User.username == current_username).first()
    if not user:
        raise HTTPException(status_code=401, detail="User profile mismatch.")
        
    return db.query(models.JobApplication).filter(models.JobApplication.user_id == user.id).order_by(models.JobApplication.id.desc()).all()

@app.put("/jobs/{job_id}", response_model=schemas.JobResponse)
def update_job(
    job_id: int, 
    job_update: schemas.JobUpdate, 
    db: Session = Depends(database.get_db),
    current_user: str = Depends(get_current_user)
):
    db_job = get_user_job_or_raise(job_id, db, current_user)
    
    for key, val in job_update.model_dump(exclude_unset=True).items():
        setattr(db_job, key, val)
        
    db.commit()
    db.refresh(db_job)
    return db_job

@app.post("/generate-assets")
def generate_assets(
    request: schemas.AssetGenerationRequest, 
    db: Session = Depends(database.get_db),
    current_user: str = Depends(get_current_user)
):
    db_job = get_user_job_or_raise(request.job_id, db, current_user)
    
    jd_payload = {
        "job_title": db_job.job_title,
        "company_name": db_job.company_name,
        "raw_jd": db_job.raw_jd
    }
    
    generated_text = ai_service.generate_tailored_assets(request.resume_base, jd_payload, request.asset_type)
    if request.asset_type == "resume":
        db_job.tailored_resume = generated_text
    else:
        db_job.tailored_cover_letter = generated_text
        
    db.commit()
    return {"status": "success", "data": generated_text}

@app.post("/jobs/upload-resume-tailor")
async def upload_resume_and_tailor(
    job_id: int = Form(...),
    asset_type: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(database.get_db),
    current_user: str = Depends(get_current_user)
):
    contents = await file.read()
    resume_text = ""

    if contents.startswith(b'\xff\xd8\xff'):
        raise HTTPException(
            status_code=400, 
            detail="The uploaded file appears to be an IMAGE (JPEG). Please upload a text document."
        )

    if file.filename.endswith(".pdf"):
        try:
            import pypdf
            pdf_reader = pypdf.PdfReader(io.BytesIO(contents))
            resume_text = "".join([page.extract_text() for page in pdf_reader.pages])
        except ImportError:
            raise HTTPException(status_code=500, detail="Missing engine library support.")
    else:
        resume_text = contents.decode("utf-8")

    db_job = get_user_job_or_raise(job_id, db, current_user)

    jd_payload = {
        "job_title": db_job.job_title,
        "company_name": db_job.company_name,
        "raw_jd": db_job.raw_jd
    }

    generated_text = ai_service.generate_tailored_assets(resume_text, jd_payload, asset_type)

    if asset_type == "resume":
        db_job.tailored_resume = generated_text
    else:
        db_job.tailored_cover_letter = generated_text

    db.commit()
    return {"status": "success", "data": generated_text}

@app.get("/jobs/{job_id}/download-pdf")
def download_job_pdf(
    job_id: int, 
    asset_type: str = "resume",
    db: Session = Depends(database.get_db),
    current_user: str = Depends(get_current_user)
):
    db_job = get_user_job_or_raise(job_id, db, current_user)
        
    text_content = db_job.tailored_resume if asset_type == "resume" else db_job.tailored_cover_letter
    if not text_content:
        raise HTTPException(status_code=404, detail=f"No tailored {asset_type} content found to download yet.")
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
    story = []
    
    styles = getSampleStyleSheet()
    normal_style = ParagraphStyle('ResumeBody', parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=6)
    
    clean_text = text_content.replace("**", "").replace("*", "")
    for line in clean_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 10))
            continue
            
        lower_line = stripped.lower()
        if (lower_line.startswith("here is") or 
            lower_line.startswith("here's") or 
            lower_line.startswith("note") or 
            lower_line.startswith("this is a tailored") or
            lower_line.startswith("i have reorganized") or
            lower_line.startswith("i've highlighted")):
            continue
            
        story.append(Paragraph(stripped, normal_style))
            
    doc.build(story)
    buffer.seek(0)
    
    filename = f"{db_job.company_name}_Tailored_{asset_type.capitalize()}.pdf".replace(" ", "_")
    return StreamingResponse(buffer, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename={filename}"})


@app.get("/jobs/{job_id}/download-docx")
def download_job_docx(
    job_id: int, 
    asset_type: str = "resume",
    db: Session = Depends(database.get_db),
    current_user: str = Depends(get_current_user)
):
    db_job = get_user_job_or_raise(job_id, db, current_user)
        
    text_content = db_job.tailored_resume if asset_type == "resume" else db_job.tailored_cover_letter
    if not text_content:
        raise HTTPException(status_code=404, detail=f"No tailored {asset_type} content found to download yet.")
    
    doc = Document()
    clean_text = text_content.replace("**", "").replace("*", "")
    
    for line in clean_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
            
        lower_line = stripped.lower()
        if (lower_line.startswith("here is") or 
            lower_line.startswith("here's") or 
            lower_line.startswith("note") or 
            lower_line.startswith("this is a tailored") or
            lower_line.startswith("i have reorganized") or
            lower_line.startswith("i've highlighted")):
            continue
            
        doc.add_paragraph(stripped)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"{db_job.company_name}_Tailored_{asset_type.capitalize()}.docx".replace(" ", "_")
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.delete("/jobs/{job_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_job(job_id: int, db: Session = Depends(database.get_db), current_user: str = Depends(get_current_user)):
    db_job = db.query(models.JobApplication).filter(models.JobApplication.id == job_id).first()
    
    if not db_job:
        raise HTTPException(
            status_code=404, 
            detail="The requested application profile target container could not be found to delete."
        )
    
    db.delete(db_job)
    db.commit()
    return None

@app.post("/register", response_model=schemas.UserResponse, status_code=201)
def register_user(user_in: schemas.UserCreate, db: Session = Depends(database.get_db)):
    existing_user = db.query(models.User).filter(
        (models.User.username == user_in.username) | (models.User.phone == user_in.phone)
    ).first()
    if existing_user:
        raise HTTPException(status_code=400, detail="Username or Phone Number parameter is already registered.")
    
    hashed = hash_password(user_in.password)
    new_user = models.User(
        username=user_in.username, 
        hashed_password=hashed,
        phone=user_in.phone
    )
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    return new_user

@app.post("/auth/forgot-password", status_code=200)
def forgot_password_reset(payload: schemas.PasswordResetRequest, db: Session = Depends(database.get_db)):
    user = db.query(models.User).filter(models.User.phone == payload.phone).first()
    if not user:
        raise HTTPException(status_code=404, detail="No registered workspace profile matches this phone identifier.")
    
    user.hashed_password = hash_password(payload.new_password)
    db.commit()
    return {"message": "Security credential credentials vector reset completed successfully."}


# ─── SECURE ADMINISTRATIVE MANAGEMENT PANEL ROUTING ENDPOINTS ───

@app.get("/admin/users")
def get_all_users(db: Session = Depends(database.get_db), current_username: str = Depends(get_current_user)):
    # 1. Fetch current user context
    user = db.query(models.User).filter(models.User.username == current_username).first()
    if not user or not getattr(user, "is_admin", False):
        raise HTTPException(status_code=403, detail="Access denied. Administrative privileges required.")
        
    # 2. Return high-level security user profiles summary
    users = db.query(models.User).all()
    return [
        {
            "id": u.id, 
            "username": u.username, 
            "phone": u.phone, 
            "is_admin": getattr(u, "is_admin", False)
        } 
        for u in users
    ]

@app.delete("/admin/users/{user_id}")
def admin_delete_user(user_id: int, db: Session = Depends(database.get_db), current_username: str = Depends(get_current_user)):
    # 1. Verify authorization context
    user = db.query(models.User).filter(models.User.username == current_username).first()
    if not user or not getattr(user, "is_admin", False):
        raise HTTPException(status_code=403, detail="Access denied. Administrative privileges required.")
        
    # 2. Verify target existence
    user_to_delete = db.query(models.User).filter(models.User.id == user_id).first()
    if not user_to_delete:
        raise HTTPException(status_code=404, detail="Target user profile record missing from system index.")
        
    # 3. Prevent admin self-lockout/destruction hazards
    if user_to_delete.id == user.id:
        raise HTTPException(status_code=400, detail="Administrative profile self-deletion is blocked.")

    # 4. Clean purge dependent job metadata rows to preserve database cascading integrity
    db.query(models.JobApplication).filter(models.JobApplication.user_id == user_id).delete()

    db.delete(user_to_delete)
    db.commit()
    return {"message": f"User account container {user_id} successfully purged from database cloud node clusters."}
    # ─── AUTO-INITIALIZE ADMINISTRATIVE USER OVERRIDES ───
@app.on_event("startup")
def create_default_admin():
    import database, models, os
    
    try:
        admin_user = os.getenv("ADMIN_USERNAME", "whitedevil")
        
        db = database.SessionLocal()
        existing = db.query(models.User).filter(models.User.username == admin_user).first()
        
        if existing:
            print(f"Purging outdated database record for: {admin_user} to allow clean state reset.")
            db.delete(existing)
            db.commit()
            
        db.close()
        print("Database admin row cleared. Ready for clean setup or frontend enrollment!")
    except Exception as e:
        print(f"Admin auto-creation warning: {e}")